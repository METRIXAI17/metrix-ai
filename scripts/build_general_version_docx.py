"""Build the general-version Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "KARIM_METRIX_GENERAL_VERSION_2026-08-25.docx"

TEAL = RGBColor(0x0F, 0x76, 0x6E)
INK = RGBColor(0x14, 0x18, 0x1C)
MUTED = RGBColor(0x4B, 0x55, 0x63)


def _set_run(run, *, size=12, bold=False, color=INK, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _set_run(run, size=18 if level == 1 else 14, bold=True, color=TEAL)
    return p


def p(doc, text, *, italic=False, size=12, space_after=8):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    _set_run(run, size=size, italic=italic)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            _set_run(run, size=12)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(item, style="List Number")
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            _set_run(run, size=12)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htxt)
        _set_run(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = cell._tePr if hasattr(cell, "_tePr") else cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0F766E")
        shd.set(qn("w:val"), "clear")
        shading.append(shd)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            _set_run(run, size=11)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    style.font.color.rgb = INK

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Karim Metrix")
    _set_run(r, size=26, bold=True, color=TEAL)
    sub = doc.add_paragraph()
    r = sub.add_run("Генеральная версия · 25 августа 2026 · релиз 1.6.0")
    _set_run(r, size=12, color=MUTED)
    p(doc, "Публичный голос: @karimmetrix  ·  билдер на столе  ·  демо как магистраль к платному.", italic=True)

    h(doc, "Вердикт", 1)
    p(
        doc,
        "Проект состоятелен как соло-продукт с узкой нишей: уникальные финансовые модели, которые садятся в чужой контур, и агенты, которые эти модели держат. Не состоятелен как «платформа всего». Не сигналы. Не гарантия доходности.",
    )
    p(
        doc,
        "До правки система была умнее коммуникации. Наружу торчали внутренние имена: «работа по запросу», «флагманские карточки», «терминал», JSON вместо вещи, которую можно сохранить. Рынок это не покупает.",
    )

    h(doc, "Что сломано было в продукте", 1)
    table(
        doc,
        ["Было", "Почему ломает"],
        [
            ["Кнопки-жаргон", "Человек не знает, что нажать."],
            ["Бот = запуск Mini App", "Нет личности Карима."],
            ["Демо = JSON и «режим»", "Нечего сохранить и переслать."],
            ["Нет петли «зашло / мимо»", "Неясно, что является товаром."],
            ["Слишком много поверхностей", "Внимание распадается."],
            ["Стратегии не названы", "Origin-story не стала объектом."],
            ["Агенты не были дверью", "А билдер на столе как раз про это."],
        ],
    )

    h(doc, "Генеральная версия", 1)
    p(
        doc,
        "Metrix — билдер на столе. Из ситуации собирается демо-артефакт. Если он зашёл — это товар: уникальная финансовая модель и/или агент, который её держит. Если не зашёл — не продаём.",
    )
    p(doc, "Магистраль:", italic=True)
    bullets(
        doc,
        [
            "Ситуация своими словами.",
            "Демо-артефакт бесплатно.",
            "Зашло / почти / мимо — майнер ценности.",
            "Пилот 14 дней: посадка именно этого артефакта.",
        ],
    )
    p(
        doc,
        "Майнер ценности — не охваты. Это артефакты, которые человек пометил «зашло». Остальное не тянем. Paid не начинается с каталога SKU. Paid начинается с резонанса.",
    )

    h(doc, "Личность и бот", 1)
    p(
        doc,
        "Молодой человек, которому что-то удалось в трейдинге. Ниша — финансовые модели как посадка в чужой проект. Посты на X иногда собирают людей. Фриланс, когда есть окно. Когда окна нет — спокойно залипает.",
    )
    table(
        doc,
        ["Было", "Стало"],
        [
            ["Открыть Metrix AI", "Открыть билдер"],
            ["Работа по запросу", "Демо"],
            ["Карточки / Флагманы", "Стратегии"],
            ["Промо / Терминал на первом экране", "Убраны. Не первая дверь."],
            ["—", "Агенты"],
            ["—", "Посты"],
        ],
    )
    p(doc, "Любой свободный текст тоже собирает демо. Это и есть магистраль, а не спрятанная кнопка.")

    h(doc, "Чем полезен билдер на столе", 1)
    p(doc, "Агент здесь не чат. Агент держит модель: что считать деньгами, когда молчать, какой артефакт отдать человеку.")

    h(doc, "B2B SaaS и IT, 50–500 человек", 2)
    p(
        doc,
        "Тонут не в коде. Тонут в решениях без экономики: фича «вроде нужна», пилот «вроде идёт», победа нигде не определена. Команда уже купила чаты. Билдер сажает в агента финмодель продуктовой линии: сырой запрос → unit-экономика фичи + условие остановки + артефакт раскатки. Агент сидит на входящих фичах, не в HR-чате.",
    )

    h(doc, "Агентства digital и performance", 2)
    p(
        doc,
        "Маржа сгорает на онбординге. Метод живёт в головах аккаунтов. Агент входа: бриф → геометрия оффера → пакет на 14 дней с цифрой маржи. Если метод нельзя посадить в агента — метода нет, есть геройство.",
    )

    h(doc, "Образовательные проекты и онлайн-школы", 2)
    p(
        doc,
        "Контент путают с деньгами. FAQ-бот по уроку школу не спасает. Агент когорты: урок или трафик → следующий платный шаг + экономика потока. Модель LTV внимания: какой кусок программы продаёт, какой только греет эго.",
    )

    h(doc, "E-commerce с высоким средним чеком (B2B и B2C)", 2)
    p(
        doc,
        "Дорогой клик. Высокий чек должен оплатить отношение, не ещё одну ставку. Общий виджет на сайте ведёт себя как для футболок. Агент заказа: покупка → цикл капитала (дожим, возврат, повтор) с разными правилами для B2B и B2C.",
    )

    h(doc, "Три торговые модели", 1)
    p(doc, "Модели мышления и журнала. Не сигналы, не оферта доходности, не робот с ордерами на бирже.", italic=True)

    h(doc, "Target Place — золото", 2)
    p(
        doc,
        "Вход и выход — места, не ощущения. Три типа мест: магнит, origin, инвалидация. Вход, когда цена пришла в заранее отмеченное место и отвергла его. Выход — противоположный магнит или конец NY, если место не уважили. Между местами — воздух.",
    )

    h(doc, "Demand — крипта", 2)
    p(
        doc,
        "Местные истории, которые стреляют в коротком окне. Сначала окно спроса (часто 24–72 часа вокруг катализатора), потом имя. Вход, когда спрос виден внутри окна. Выход — конец окна, даже если «могло ещё». Идентичность к монете убивает больше волатильности.",
    )

    h(doc, "Ampli — Америка", 2)
    p(
        doc,
        "Сборщик амплитуды, не предсказатель направления. Сжатие → расширение диапазона в cash-сессию → съём, когда амплитуда сдохла. Сторона берётся по факту расширения, не по мнению.",
    )

    h(doc, "Посты для X", 1)
    p(
        doc,
        "14 черновиков лежат в docs/x-posts-karimmetrix-2026-08-25.md. Картинки без текста: target-place-gold, demand-crypto, ampli-us, desktop-builder. Ритм на две недели: модель / ниша / стратегия / демо-магистраль / личность. Не каждый день. Не треды.",
    )

    h(doc, "Что внедрено в код", 1)
    bullets(
        doc,
        [
            "Движки: demo_highway, strategies, agent_studio, resonance, x_posts, voice.",
            "Telegram-бот: личность, понятные кнопки, демо из свободного текста, петля зашло/почти/мимо.",
            "Mini App: четыре вкладки — Главная, Демо, Стратегии, Агенты. Артефакт вместо JSON.",
            "API: /api/v1/miniapp/demo, resonate, strategy, agent, posts.",
            "Релиз 1.6.0. Бот стартует рядом с API, если задан TELEGRAM_BOT_TOKEN.",
        ],
    )

    h(doc, "Оценка после правки", 1)
    p(doc, "Можно показывать людям. Нельзя обещать доходность, scale и «агент сам всё сделает».")
    bullets(
        doc,
        [
            "Origin story совпадает с оффером.",
            "Демо наконец является объектом.",
            "Резонанс даёт честный SKU: продаём то, что зашло.",
            "Ядро детерминированное — для части людей это будет «умный шаблон». Не притворяться магией.",
            "Пилот 14 дней — ручная работа, пока нет очереди.",
            "Оплата в боте выключена. Магистраль ведёт к пилоту руками — честнее мёртвых счетов.",
        ],
    )

    h(doc, "Дальнейшие шаги", 1)
    numbered(
        doc,
        [
            "Опубликовать пять постов с картинками (модель, агент, Target Place, Demand, демо-магистраль). Смотреть, что срезонировало.",
            "Прогнать бота живьём: /start, Демо, три стратегии, четыре ниши, зашло/мимо. Поправить фразы, которые режут глаз.",
            "Десять живых демок знакомым из четырёх ниш. Считать только «зашло».",
            "Один пилот за деньги с артефактом, который уже зашёл. Не продавать пакет «вообще Metrix».",
            "Не включать оплату в боте, пока не будет трёх пилотов руками.",
            "Следующим релизом сузить публичный сайт до той же магистрали. В этом деплое узким местом был бот.",
            "Журнал сделок оставить тихой функцией, не витриной.",
            "Не подключать биржевые API. Модель — не робот. Робот — отдельный продукт, риск и юридический контур.",
        ],
    )

    h(doc, "Деплой", 1)
    p(doc, "GitHub METRIXAI17/metrix-ai, ветка main. Railway — API и бот. Vercel — public/, Mini App на /tg/. После выкладки: health API, открыть /tg/, нажать /start в Telegram.")

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(24)
    r = footer.add_run("Karim Metrix  ·  не сигналы  ·  если зашло — это товар")
    _set_run(r, size=10, italic=True, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
