#!/usr/bin/env python3
"""Build Metrix AI 12-month financial operating plan DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "METRIX_AI_12M_FINANCIAL_OPERATING_PLAN.docx"


def set_run(run, bold=False, size=10, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor(*color)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=True, size=16, color=(15, 23, 42))
    p.paragraph_format.space_after = Pt(8)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=True, size=12, color=(13, 148, 136))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=True, size=11, color=(30, 64, 175))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold, size=size, color=(30, 41, 59))
    r.italic = italic
    p.paragraph_format.space_after = Pt(4)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    r = p.add_run(text)
    set_run(r, size=10, color=(30, 41, 59))


def shade_cell(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    )


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=8, color=(255, 255, 255))
        shade_cell(cell, "0F766E")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_run(r, bold=(ci == 0), size=8, color=(15, 23, 42))
            if ri % 2 == 1:
                shade_cell(cell, "F0FDFA")
    doc.add_paragraph()
    return table


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    h1(doc, "Metrix AI — 12-Month Financial & Go-to-Market Operating Plan")
    para(
        doc,
        "Confidential planning model · Karim Metrix / Market Units · 21 July 2026",
        italic=True,
        size=9,
    )
    para(
        doc,
        "Assumption mode: decisions are taken; Stage-2 pipeline is complete; narrative layer and "
        "paid package are finished in Months 1–2; business operates 12 months under a hard capital "
        "envelope of USD 130,000. Figures are realistic planning estimates (not audited). Prior "
        "models used only information-generation potential; this plan binds numbers to a real "
        "commercial funnel (demos → clients → reorders) and to Sell-Ops as Objectly access + OpeningEdge.",
    )

    h2(doc, "0. Executive Snapshot (Base Case)")
    add_table(
        doc,
        ["Metric", "Value", "Note"],
        [
            ["Investment envelope (12 mo)", "USD 130,000", "Hard cap — fully allocated"],
            ["Year-1 gross revenue (base)", "USD 224,000", "100 new + 60 reorder txs"],
            ["Year-1 free cash (planning)", "~USD 72,000", "After fees & contingency"],
            ["Demo capacity", "400 / year", "Free + light oriented demos"],
            ["Paying clients (new)", "100", "25% demo→client (aggressive)"],
            ["Repeat buyers", "30 × 2 orders", "60 reorder events"],
            ["Blended first-order ARPU", "USD 1,130", "$290 / $1,490 / $2,490 / $1,790 mix"],
            ["Core cloud ops (12 mo)", "USD 3,600", "VPS + backups + domain + monitor"],
            ["Avoided pure-LLM API bill", "~USD 100,000", "Counterfactual multi-agent stack"],
            ["Market share SOM Y1", "0.02–0.08% of SAM", "See §1"],
            ["Business life horizon", "7–12 years", "If Objectly/Edge expand; else 3–5"],
        ],
    )

    h2(doc, "1. Market Size, Share & Longevity")
    para(
        doc,
        "TAM/SAM/SOM frame productized operational orientation + pilot packaging for AI agencies, "
        "cloud/FinOps, chip, telecom, brand ops, and device assembly — not general chatbot SaaS.",
    )
    h3(doc, "1.1 Market frames (USD annual)")
    add_table(
        doc,
        ["Layer", "Definition", "USD estimate", "Metrix claim Y1–Y3"],
        [
            ["TAM", "Global digital ops / AI delivery consulting + productized diagnostics", "35–55B", "Negligible"],
            ["SAM", "EN/RU mid-market buyers of orientation/pilot packs (6 verticals)", "1.8–3.2B", "Tiny"],
            ["SOM Y1", "Captured under 1 founder + lean team + $130k", "0.15–0.35M", "0.02–0.08% SAM"],
            ["SOM Y3", "With Objectly sell-ops + OpeningEdge + Harness", "0.8–2.0M", "0.05–0.15% SAM"],
        ],
    )
    h3(doc, "1.2 Extractable money & lifetime")
    add_table(
        doc,
        ["Horizon", "Conservative", "Base (plan)", "Upside", "Driver"],
        [
            ["Y1 revenue", "95,000", "224,000", "380,000", "Funnel + narrative pack"],
            ["Y2 revenue", "180,000", "420,000", "750,000", "Reorders + Objectly"],
            ["Y3 revenue", "280,000", "700,000", "1,400,000", "Harness + Data Market"],
            ["Founder extract Y1", "24–36k salary", "36k salary + 10–25k profit", "Higher profit", "Stay in $130k"],
            ["Business lifetime", "5–7 yrs", "7–12 yrs", "12–20 yrs", "Access/asset layer"],
        ],
    )
    para(
        doc,
        "Longevity: pure “generate new information” products decay in 3–5 years as LLM wrappers "
        "commoditize. Metrix extends life via (1) proprietary pipeline + QC, (2) Objectly standardized "
        "virtual assets of clients-of-clients, (3) OpeningEdge for end consumers, (4) sell-ops access rights.",
    )

    h2(doc, "2. Commercial Funnel (Decided Operating Point)")
    h3(doc, "2.1 Throughput")
    add_table(
        doc,
        ["Stage", "Volume / year", "Rate", "Comment"],
        [
            ["Demos (oriented runs)", "400", "—", "X + ads + outbound"],
            ["New paying clients", "100", "25% of demos", "Aggressive vs 8–15% median"],
            ["Repeat buyers", "30 of 100", "30%", "Healthy for pilots"],
            ["Reorder events", "60", "30 × 2", "Deeper package / 2nd product"],
            ["Total paid transactions", "160", "100 + 60", "Primary P&L unit"],
        ],
    )
    h3(doc, "2.2 Price book")
    add_table(
        doc,
        ["SKU", "USD", "Role"],
        [
            ["Orientation + levers", "290", "Entry diagnose"],
            ["Pilot 14–30d", "1,490", "Core cash engine"],
            ["Full Package", "2,490", "Packed narrative product"],
            ["Additional capability", "1,790", "Objectly / Edge / specialty"],
            ["Cloud FinOps board", "1,890", "Cloud vertical flagship"],
        ],
    )
    h3(doc, "2.3 First-order mix (100 new clients)")
    add_table(
        doc,
        ["SKU", "Clients", "Revenue USD"],
        [
            ["Orientation $290", "45", "13,050"],
            ["Pilot $1,490", "30", "44,700"],
            ["Full Package $2,490", "15", "37,350"],
            ["Capability $1,790", "10", "17,900"],
            ["TOTAL first orders", "100", "113,000"],
        ],
    )
    h3(doc, "2.4 Reorders (30 clients × 2)")
    para(
        doc,
        "Reorder mix assumption: 50% Pilot · 30% Full · 20% Capability → blended USD 1,850 / reorder.",
    )
    add_table(
        doc,
        ["Item", "Calc", "USD"],
        [
            ["Reorder events", "60", "—"],
            ["Blended reorder price", "1,850", "—"],
            ["Reorder revenue", "60 × 1,850", "111,000"],
            ["YEAR-1 GROSS REVENUE", "113k + 111k", "224,000"],
        ],
    )
    h3(doc, "2.5 Conservative alternate (12% conversion)")
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Paying clients from 400 demos", "48"],
            ["First-order revenue @ ARPU 1,130", "~54,000"],
            ["Reorders (approx)", "~31,000"],
            ["Year-1 gross (conservative)", "~85,000"],
        ],
    )
    para(doc, "Plan uses BASE ($224k). Capital of $130k must survive near the conservative floor.")

    h2(doc, "3. Cost Structure — Full 12-Month Budget (USD 130,000)")
    para(
        doc,
        "Complete capital allocation as if all decisions are made. Includes items founders often miss "
        "(legal, tax/PSP reserve, tools, contingency, contractor float).",
    )
    h3(doc, "3.1 Master budget (fits $130,000)")
    add_table(
        doc,
        ["#", "Category", "USD", "%", "What it buys"],
        [
            ["1", "Founder salary (12 mo)", "36,000", "27.7%", "$3,000/mo lean living wage"],
            ["2", "Narrative + paid-pack expert", "18,000", "13.8%", "Conclusions system + Full Package"],
            ["3", "Advertising & paid acquisition", "22,000", "16.9%", "X + LinkedIn/search + creative"],
            ["4", "Content & X organic engine", "9,600", "7.4%", "Visuals, editor, tools, collabs"],
            ["5", "Core Operational Cloud", "3,600", "2.8%", "VPS, backups, domain, CDN, monitoring"],
            ["6", "Optional LLM polish (not core)", "2,400", "1.8%", "Copy polish only"],
            ["7", "R&D software + QC / consistency", "12,000", "9.2%", "2nd product sales-support QC"],
            ["8", "Harness / auto sell-ops integrate", "6,000", "4.6%", "API, tracking, CRM hooks"],
            ["9", "Objectly + OpeningEdge MVP pack", "5,000", "3.8%", "Separate product scaffolds"],
            ["10", "Legal / contracts / privacy", "3,500", "2.7%", "Offers, ToS, DPA"],
            ["11", "Accounting / tax / PSP reserve", "4,500", "3.5%", "Books + ~3% payment fees float"],
            ["12", "Software tools stack", "2,400", "1.8%", "Design, analytics, email, PM"],
            ["13", "Hardware / coworking / misc", "1,800", "1.4%", "Buffer ops"],
            ["14", "Contingency / runway buffer", "3,200", "2.5%", "Missed CAC, rework, sick time"],
            ["", "TOTAL", "130,000", "100%", "Hard envelope"],
        ],
    )

    h3(doc, "3.2 Core Operational Cloud vs AI API (~$100k)")
    add_table(
        doc,
        ["Architecture", "12-mo USD", "Scales with", "QC / consistency"],
        [
            ["A. Pure multi-agent LLM API", "80,000–120,000", "Tokens × demos × retries", "High variance"],
            ["B. Hybrid rules + LLM", "25,000–45,000", "Still token-heavy funnel", "Medium"],
            ["C. Metrix Core Ops Cloud (chosen)", "2,400–4,800", "VPS + storage + light polish", "Deterministic spine"],
            ["Plan C + polish line", "6,000", "Cloud 3.6k + polish 2.4k", "Matches budget"],
            ["Avoided vs $100k API", "~94,000", "Freed for people & growth", "Strategic moat"],
        ],
    )
    para(
        doc,
        "The ~$100,000 is a counterfactual annual AI-API burn for 400 demos + deep agent runs + free-preview "
        "tax. Metrix replaces it with ~$3.6k cloud + $2.4k polish. That ~$94k is not automatic profit — "
        "it is capital reallocated to salary, ads, and narrative expertise.",
    )

    h3(doc, "3.3 R&D + QC (2nd product — sales support)")
    add_table(
        doc,
        ["Workstream", "USD", "Output"],
        [
            ["Anti-generic / consistency filters", "3,500", "No empty repeated conclusions"],
            ["Eval harness (golden cases × 6 industries)", "3,000", "Narrative regression tests"],
            ["Guiding-question engine upgrades", "2,500", "Stronger intake before re-run"],
            ["Sales-support second product MVP", "3,000", "Reseller/Harness QC playbooks"],
            ["TOTAL R&D + QC", "12,000", "Not buyable cleanly from raw LLM"],
        ],
    )

    h3(doc, "3.4 Advertising lock")
    add_table(
        doc,
        ["Channel", "USD/yr", "Role", "Rough CAC"],
        [
            ["X ads + boosts", "8,000", "Authority + demos", "$40–120 / demo"],
            ["LinkedIn / niche B2B", "7,000", "Agency / FinOps buyers", "$80–200 / demo"],
            ["Search / retargeting", "4,000", "Bottom funnel", "$30–90 / demo"],
            ["Ad creative production", "3,000", "Static + short video", "—"],
            ["TOTAL PAID", "22,000", "Primary paid acquisition", "Target CAC ≤ $220 / client"],
        ],
    )
    para(
        doc,
        "If $22k fully attributed to 100 clients → media CAC $220. Healthy when LTV ≥ $1,800–2,500 "
        "(true once reorders land). Organic X is separate ($9.6k).",
    )

    h3(doc, "3.5 Content & X organic")
    add_table(
        doc,
        ["Item", "USD", "Cadence"],
        [
            ["Founder time (inside salary)", "0 extra", "4–6 posts/week"],
            ["Visuals / short clips contractor", "4,800", "8–12 assets/mo"],
            ["EN micro-editor for posts", "2,400", "2–4 hrs/week"],
            ["Tools (scheduler, analytics)", "1,200", "Annual"],
            ["Community / collabs", "1,200", "Quarterly"],
            ["TOTAL content & X", "9,600", "Supports 400 demos"],
        ],
    )

    h3(doc, "3.6 Expert narrative & packed paid product")
    add_table(
        doc,
        ["Deliverable", "USD", "Timing"],
        [
            ["Conclusion / reading system design", "6,000", "M1–M2"],
            ["Full Package narrative pack (7 types)", "5,000", "M2–M3"],
            ["Guiding-question trees (6 industries)", "3,000", "M2–M4"],
            ["Ongoing narrative QA (8 × $500)", "4,000", "M3–M12"],
            ["TOTAL expert narrative", "18,000", "Critical path after Stage-2"],
        ],
    )
    para(
        doc,
        "Bridge from “oriented template” to sellable long analysis. Pipeline is strong; without this "
        "spend, paid package quality remains the weak point shown in tests.",
    )

    h3(doc, "3.7 Harness automated sales — cost & popularity")
    add_table(
        doc,
        ["Item", "USD / note", "Assessment"],
        [
            ["Integration + CRM hooks Y1", "6,000", "Build + iterate"],
            ["Popularity Y1", "Low–medium", "Niche operators, not mass"],
            ["Popularity Y2–Y3", "Medium", "If Objectly becomes sell-ops primitive"],
            ["Y1 pure auto-sales revenue", "5–20k", "Do not depend on this"],
            ["Role", "Force multiplier", "After narrative product is solid"],
        ],
    )

    h2(doc, "4. Team & Founder Compensation")
    h3(doc, "4.1 Founder salary")
    add_table(
        doc,
        ["Option", "Monthly", "Annual", "Fit $130k"],
        [
            ["Survival lean (LOCKED)", "3,000", "36,000", "Yes"],
            ["Market founder draw", "5,000–7,000", "60–84k", "Breaks envelope"],
            ["Deferred + profit share", "2,000 + bonus", "24k + upside", "If high risk tolerance"],
        ],
    )
    para(
        doc,
        "Locked: USD 3,000/mo from capital for 12 months. From Month 7+, if base revenue hits, raise "
        "draw from profit without new investment.",
    )
    h3(doc, "4.2 Who else is needed")
    add_table(
        doc,
        ["Role", "Mode", "Cost in plan", "When"],
        [
            ["Founder (product, sales, X, architecture)", "FT", "Salary $36k", "Always"],
            ["Narrative / product copy expert", "Contractor", "$18k", "M1–M12"],
            ["Content visualist", "PT contractor", "Inside $9.6k", "M1–M12"],
            ["EN micro-editor", "PT", "Inside content/expert", "M1–M12"],
            ["Closer / BDR (optional H2)", "Rev-share 8–12%", "From revenue", "If >40 clients"],
            ["Full-time engineer Y1", "None", "$0", "Hire Y2 if Objectly scales"],
            ["Accountant (outsourced)", "Monthly", "Inside $4.5k", "M1"],
        ],
    )

    h2(doc, "5. Sell-Ops = Objectly Access + OpeningEdge")
    para(
        doc,
        "Sell-ops is defined as access to an Object — a standardized virtual asset of the businesses of "
        "your clients’ clients — and to end consumers who adopt OpeningEdge. Prior models counted only "
        "new-information generation; this plan adds access-rights economics.",
    )
    add_table(
        doc,
        ["Layer", "What is sold", "Y1 weight", "Y2–Y3 role"],
        [
            ["Core packs", "Orient / pilot / full", "85–90%", "Still core"],
            ["Objectly access", "Standardized VA of client businesses", "5–10%", "Sell-ops backbone"],
            ["OpeningEdge", "Edge/archetype layer for end users", "Small Y1", "Consumer wedge"],
            ["Data Market", "Secondary insight/asset exchange", "Scaffold", "Take-rate later"],
            ["Harness automation", "Low-touch distribution", "Indirect", "Volume lever"],
        ],
    )
    add_table(
        doc,
        ["Objectly unit (planning)", "Price band USD", "Y1 attach"],
        [
            ["Access / object metadata pack", "400–1,790", "10–20% of pilots"],
            ["Reseller portfolio of objects", "2,000–5,000", "Few lighthouse deals"],
        ],
    )

    h2(doc, "6. Savings & Average Forecast KPIs")
    h3(doc, "6.1 Architecture savings (plan volume)")
    add_table(
        doc,
        ["KPI", "Pure LLM stack", "Metrix plan", "Delta"],
        [
            ["AI/API variable cost", "~100,000", "2,400 polish", "~97,600 saved"],
            ["Core cloud", "5–10k+", "3,600", "Lower"],
            ["QC / consistency cost", "High (human+tokens)", "12,000 R&D QC", "Controllable"],
            ["Ops cost / demo", "15–40", "1–3", "Order of magnitude"],
            ["Ops cost / paid client", "60–150", "5–15", "Before CAC"],
        ],
    )
    h3(doc, "6.2 Business KPI targets Y1")
    add_table(
        doc,
        ["KPI", "Target", "Note"],
        [
            ["Demo → client", "15–25%", "Plan uses 25%; manage ≥15%"],
            ["Gross revenue", "224,000", "§2"],
            ["First ARPU", "1,130", "Mix-driven"],
            ["LTV 12–18 mo", "2,000–2,800", "First + reorders"],
            ["CAC blended", "150–280", "Ads + organic"],
            ["LTV:CAC", "≥ 8:1", "Services-software hybrid"],
            ["Gross margin after delivery", "55–70%", "Founder-heavy Y1"],
            ["Runway on $130k alone", "12 months", "By design"],
            ["Cash break-even month", "M6–M8", "If base funnel hits"],
        ],
    )
    h3(doc, "6.3 Simplified P&L sketch (base)")
    add_table(
        doc,
        ["Line", "USD"],
        [
            ["Gross revenue", "224,000"],
            ["Payment fees ~3%", "(6,700)"],
            ["Net receipts", "217,300"],
            ["Capital deployed as opex*", "(130,000)"],
            ["Apparent surplus vs capital", "87,300"],
            ["Delivery overload contingency", "(15,000)"],
            ["Planning free cash Y1", "~72,000"],
        ],
    )
    para(
        doc,
        "* $130k is spent as budgeted opex. Revenue builds a separate cash pile for Y2 — do not "
        "double-count salary as both capital and COGS without care.",
    )

    h2(doc, "7. When $130k Is Enough / Breaks")
    add_table(
        doc,
        ["Condition", "Result"],
        [
            ["Base funnel (100 + reorders)", "Sufficient; cash for Y2"],
            ["Conservative (~48 clients)", "Survive; cut ads $6–8k if needed"],
            ["Founder $5k/mo without revenue", "Breaks by M9–M10"],
            ["Pure LLM path", "Needs +$80–100k or kills demos"],
            ["No narrative expert", "Pipeline OK; paid pack stays weak"],
            ["Harness expected to print money Y1", "Fails — experiment only"],
        ],
    )

    h2(doc, "8. Quarterly Cash Phasing")
    add_table(
        doc,
        ["Quarter", "Focus", "USD", "Milestone"],
        [
            ["Q1", "Narrative + QC + cloud + salary + content", "42,000", "Paid pack v1; questions live"],
            ["Q2", "Ads ramp + salary + Harness hooks", "36,000", "100 demos; 25 clients"],
            ["Q3", "Ads + salary + Objectly/Edge pack", "30,000", "Reorders; 60 clients cum."],
            ["Q4", "Salary + content + buffer + residual ads", "22,000", "100 clients; 30 repeat"],
            ["Total", "", "130,000", "Narrative business operating"],
        ],
    )

    h2(doc, "9. Risks Priced In")
    add_table(
        doc,
        ["Risk", "Mitigation in budget"],
        [
            ["25% conversion too high", "Conservative scenario + contingency + ad flexibility"],
            ["English / marketing gap", "$18k expert + $9.6k content/X"],
            ["Founder delivery burnout", "Cap Full Packages; rev-share closer H2"],
            ["LLM cost surprise", "Architecture C; polish capped $2.4k"],
            ["Harness hype", "Only $6k; no revenue dependency"],
            ["Old model had no real clients", "Forces 400 / 100 / 30 targets"],
        ],
    )

    h2(doc, "10. Locked Decision Summary")
    bullet(doc, "Capital: USD 130,000 for 12 months — fully allocated (§3.1).")
    bullet(doc, "Architecture: Core Operational Cloud ~$3.6k + polish $2.4k instead of ~$100k AI API.")
    bullet(doc, "People: Founder $3k/mo + narrative expert $18k + light content; no FT eng Y1.")
    bullet(doc, "Growth: Ads $22k + organic X $9.6k → support 400 demos.")
    bullet(doc, "Product priority: narrative/paid pack ($18k) + QC ($12k) after Stage-2 pipeline.")
    bullet(doc, "Sell-ops future: Objectly access + OpeningEdge; Data Market later.")
    bullet(doc, "Base Y1 revenue: USD 224,000; conservative floor ~USD 85,000.")
    bullet(doc, "Longevity: 7–12 years with asset/access layer; 3–5 if stuck as one-off reports.")

    para(doc, "")
    para(
        doc,
        "Disclaimer: Internal planning model only. Not investment advice, not a guarantee of revenue, "
        "not an audited statement. Replace assumptions with live invoices and CRM actuals.",
        italic=True,
        size=9,
    )
    para(
        doc,
        "Document version: 2026-07-21 · Metrix AI / Karim Metrix · Market Units planning.",
        italic=True,
        size=9,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
